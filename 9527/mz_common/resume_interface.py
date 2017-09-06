#! /usr/bin/evn python
# -*- coding:utf-8 -*-
import datetime
from django.http import HttpResponse
from django.shortcuts import render
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

import db.api.common.resume
from mz_common.common_interface import get_city_name
from utils.logger import logger as log
from docx import Document
from io import BytesIO
from docx.oxml.ns import qn
from docx.text.run import Run


def get_age_by_birthday(birthday):
    if isinstance(birthday, str):
        birthday = datetime.datetime.strptime(birthday, "%Y-%m-%d")
    today = datetime.date.today()
    age = today.year - birthday.year if today.year > birthday.year else 0
    return age


def get_work_years(start_work_time):
    """
    根据工作经历列表获取工作年限
    :param start_work_time: int ex:2012  开始工作的年份
    :return:
    """
    if start_work_time and isinstance(start_work_time, int):
        work_years = datetime.date.today().year - start_work_time
    else:
        work_years = 0
    return work_years


def get_resume_info_by_user_id(request, user_id):
    # 个人信息
    APIResult = db.api.common.resume.get_user_info(user_id)
    if APIResult.is_error():
        log.warn("get_user_info is error")
        return render(request, "404.html", dict(message=u"get_user_info is error"))
    user_info = APIResult.result()
    user_id = user_info.get("user_id", 0)
    if user_info:
        # 计算年龄
        user_info["age"] = get_age_by_birthday(user_info["birthday"])
        # 获取城市名称
        city_info = get_city_name(request, user_info.get("city_id"))
        user_info["city_name"] = city_info.get("city_name", "")
        user_info["province_name"] = city_info.get("province_name", "")
        # 计算工作年限
        try:
            user_info["work_years"] = get_work_years(user_info.get("start_work_time", 0))
        except TypeError as e:
            log.warn("user_info['start_work_time'] is None. error:%s" % e)
            user_info["work_years"] = 0

    # 工作经历
    APIResult = db.api.common.resume.list_user_work(user_id)
    if APIResult.is_error():
        log.warn("list_user_word is error.")
        return render(request, "404.html", dict(message=u"list_user_word is error."))
    user_works = APIResult.result()

    # 教育背景
    APIResult = db.api.common.resume.list_user_edu(user_id)
    if APIResult.is_error():
        log.warn("list_user_edu is error.")
        return render(request, "404.html", dict(message=u"list_user_edu is error."))
    user_edus = APIResult.result()

    return user_info, user_edus, user_works


def write_resume_to_docx(user_info, user_edus, user_works):
    """
    将简历信息写入到word
    :param user_info: dict()
    :param user_edus: list()
    :param user_works: list()
    :return:
    """
    doc = DocxExport()
    para_style = doc.set_paragraph_format()
    paragraph = doc.document.add_heading(level=0)
    run = paragraph.add_run(u'个人简历')
    doc.set_font_style(run)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置居中
    write_user_info(doc, para_style, user_info)
    write_user_work(doc, para_style, user_works)
    write_user_edu(doc, para_style, user_edus)
    file_name = "%s_%s_%s" % (user_info.get("real_name"), user_info.get("to_industry"), user_info.get("city"))
    doc.set_file_name(file_name)
    response = doc.docx_response()
    return response


def write_user_info(doc, style, user_info):
    """
    将个人信息写入到word
    :param document:
    :param user_info: dict()
    :return:
    """
    try:
        assert isinstance(user_info, dict)
        base_info = u"{0}     {1}年工作经验    |    {2}    |    {3}岁    |    现居 {4}，{5}     |  {6}".format(

            user_info.get("real_name", ""),
            user_info.get("work_years", 0),
            user_info.get("gender_name", "男"),
            user_info.get("age", 20),
            user_info.get("province_name", ""),
            user_info.get("city_name", ""),
            user_info.get("mobile", ""))
        run = doc.add_resume_head_(u"基本信息")
        doc.set_font_style(run)
        p = doc.document.add_paragraph(style=style)
        run = p.add_run(base_info)
        doc.set_font_style(run)
    except AssertionError:
        log.warn("user_info must be a dict.")


def write_user_work(doc, style, user_works):
    """
    将工作经历写入到word
    :param document:
    :param user_works:list()
    :return:
    """
    try:
        assert isinstance(user_works, list)
        run = doc.add_resume_head_(u"工作经历")
        doc.set_font_style(run)
        for work in user_works:
            work_info = u"{0}——{1} | {2}".format(
                work.get("start_time"),
                work.get("end_time"),
                work.get("company"))
            p = doc.document.add_paragraph(style=style)
            run = p.add_run(work_info)
            run.bold = True
            doc.set_font_style(run)
            work_title = u"工作岗位： {0}".format(work.get("title"))
            p = doc.document.add_paragraph(style=style)
            run = p.add_run(work_title)
            doc.set_font_style(run)
            work_content = u"工作内容： {0}".format(work.get("content").replace("\r\n", "\n"))
            p = doc.document.add_paragraph(style=style)
            run = p.add_run(work_content)
            doc.set_font_style(run)
    except AssertionError:
        log.warn("user_works must be a list.")


def write_user_edu(doc, style, user_edus):
    """
    将教育经历写入到word
    :param document:
    :param user_edus: list()
    :return:
    """
    try:
        assert isinstance(user_edus, list)
        run = doc.add_resume_head_(u"教育背景")
        doc.set_font_style(run)
        for edu in user_edus:
            edu_info = u"{0}——{1} | {2}".format(edu.get("start_time"),
                                                edu.get("end_time"),
                                                edu.get("school"))
            p = doc.document.add_paragraph(style=style)
            run = p.add_run(edu_info)
            run.bold = True
            doc.set_font_style(run)
            edu_major_title = u"专业名称： {0}         学历： {1}".format(edu.get("major"),
                                                                  edu.get("title"))
            p = doc.document.add_paragraph(style=style)
            run = p.add_run(edu_major_title)
            doc.set_font_style(run)
    except AssertionError:
        log.warn("user_edus must be a list.")


class DocxExport(object):
    """
    word（docx格式）导出
    """

    def __init__(self):
        self.document = Document()
        self._file_name = "学员简历"

    def write_bio(self):
        bio = BytesIO()
        self.document.save(bio)
        return bio

    def docx_response(self):
        """
        导出word页面
        :return: HttpResponse
        """
        bio = self.write_bio()
        response = HttpResponse(bio.getvalue(), content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=%s.docx' % self._file_name
        response['Content-Length'] = len(bio.getvalue())
        return response

    def set_file_name(self, new_file_name):
        if isinstance(new_file_name, (str, unicode)):
            if isinstance(new_file_name, unicode):
                new_file_name = new_file_name.encode("utf-8")
            new_file_name = self.file_name_format(new_file_name)
            self._file_name = new_file_name

    def set_paragraph_format(self):
        """
        设置段落样式 type=1为段落,type=2为标题
        设置['List']的样式为1.5倍行距
        :return: style
        """
        styles = self.document.styles
        style = styles['List']
        para_format = style.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_format.first_line_indent = Inches(0)
        return style

    def add_resume_head_(self, text=''):
        para = self.document.add_paragraph()
        para_format = para.paragraph_format
        para_format.space_before = WD_LINE_SPACING.DOUBLE
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(55, 93, 154)
        return run

    @staticmethod
    def set_font_style(run, font_name=u'Microsoft YaHei'):
        font_name = font_name
        try:
            assert isinstance(run, Run)
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            log.warn(u"set font style failed.info:%s" % e)

    @staticmethod
    def file_name_format(file_name):
        mark_tuple = (',', '，', '、', '。', ' ',)
        try:
            assert isinstance(file_name, (str, unicode))
            for mark in mark_tuple:
                if mark in file_name:
                    file_name = file_name.replace(mark, '_')
        except AssertionError:
            log.warn("file name must be a string.")
        return file_name
